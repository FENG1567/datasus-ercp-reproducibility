from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import re
import shutil
import textwrap
from datetime import datetime, timezone
from pathlib import Path

import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.patches import Circle, FancyBboxPatch
from matplotlib.lines import Line2D
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "manuscript" / "manuscript_content_final_blinded_v1.md"
SOURCE_SUPP_MD = ROOT / "manuscript" / "supplement_content_final_blinded_v1.md"
SOURCE_DATA = ROOT / "data" / "source_data"
SOURCE_RENDERED = ROOT / "results" / "figures"
SOURCE_PREVIEWS = SOURCE_RENDERED
TARGET = ROOT / "reproduced" / "editor_figures"
FIGURES = TARGET / "Figures"
QA = TARGET / "_qa"

MAIN_MD = TARGET / "manuscript_for_editor_blinded_v1.md"
SUPP_MD = TARGET / "supplementary_material_for_editor_v1.md"
MAIN_DOCX = TARGET / "manuscript_for_editor_blinded_v1.docx"
SUPP_DOCX = TARGET / "supplementary_material_for_editor_v1.docx"
TABLE_DOCX = TARGET / "table.docx"


PALETTE = {
    "blue": "#2C7FB8",
    "blue_soft": "#EAF3F9",
    "teal": "#2A9D8F",
    "teal_soft": "#E7F5F2",
    "orange": "#E28E2C",
    "orange_soft": "#FFF4E4",
    "red": "#C94C4C",
    "red_soft": "#FBECEC",
    "grey": "#6B7280",
    "grey_soft": "#F2F4F7",
    "line": "#D6DCE3",
    "dark": "#1F2937",
    "white": "#FFFFFF",
}


mpl.rcParams.update(
    {
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans", "sans-serif"],
        "svg.fonttype": "none",
        "pdf.fonttype": 42,
        "font.size": 7,
        "axes.spines.right": False,
        "axes.spines.top": False,
        "axes.linewidth": 0.8,
        "legend.frameon": False,
        "figure.facecolor": "white",
        "axes.facecolor": "white",
    }
)


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def ensure_empty_target() -> None:
    TARGET.mkdir(parents=True, exist_ok=True)
    existing = [p for p in TARGET.iterdir() if p.name != "desktop.ini"]
    if existing:
        raise RuntimeError(
            "Target folder is not empty; refusing to overwrite: "
            + ", ".join(p.name for p in existing[:8])
        )
    FIGURES.mkdir()
    QA.mkdir()


def read_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(SOURCE_DATA / name, dtype=str, keep_default_na=True)


def panel_letter(ax, letter: str) -> None:
    ax.text(
        -0.08,
        1.05,
        letter,
        transform=ax.transAxes,
        fontsize=9,
        fontweight="bold",
        va="bottom",
        ha="left",
        color="black",
    )


def clean_axis(ax) -> None:
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.tick_params(width=0.7, length=3, labelsize=6.5)


def rounded_card(
    ax,
    xy: tuple[float, float],
    width: float,
    height: float,
    *,
    face: str,
    edge: str,
    linewidth: float = 0.9,
    radius: float = 0.035,
) -> FancyBboxPatch:
    patch = FancyBboxPatch(
        xy,
        width,
        height,
        boxstyle=f"round,pad=0.012,rounding_size={radius}",
        facecolor=face,
        edgecolor=edge,
        linewidth=linewidth,
        transform=ax.transAxes,
        clip_on=False,
    )
    ax.add_patch(patch)
    return patch


def chip(ax, x: float, y: float, text: str, *, face: str, color: str, width: float) -> None:
    rounded_card(ax, (x, y), width, 0.095, face=face, edge=face, linewidth=0.0, radius=0.045)
    ax.text(
        x + width / 2,
        y + 0.047,
        text,
        transform=ax.transAxes,
        ha="center",
        va="center",
        fontsize=6.2,
        fontweight="bold",
        color=color,
    )


def save_figure(fig, stem: str, width_mm: float, height_mm: float) -> None:
    fig.set_size_inches(width_mm / 25.4, height_mm / 25.4)
    fig.savefig(FIGURES / f"{stem}.svg", bbox_inches="tight", facecolor="white")
    fig.savefig(FIGURES / f"{stem}.pdf", bbox_inches="tight", facecolor="white")
    fig.savefig(FIGURES / f"{stem}.png", dpi=300, bbox_inches="tight", facecolor="white")
    fig.savefig(
        FIGURES / f"{stem}.tiff",
        dpi=600,
        bbox_inches="tight",
        facecolor="white",
        pil_kwargs={"compression": "tiff_lzw"},
    )
    plt.close(fig)


def style_title(fig, title: str, subtitle: str | None = None) -> None:
    fig.suptitle(title, x=0.02, y=0.995, ha="left", va="top", fontsize=10, fontweight="bold")
    if subtitle:
        fig.text(0.02, 0.958, subtitle, ha="left", va="top", fontsize=7.2, color=PALETTE["dark"])


def build_figure_1() -> None:
    policy = read_csv("figure_1_policy_timeline_source_data.csv")
    funnel = read_csv("figure_1_cohort_funnel_source_data.csv")
    fig = plt.figure(constrained_layout=False)
    gs = fig.add_gridspec(2, 2, height_ratios=[1.45, 1.0], width_ratios=[1.48, 1.0], hspace=0.34, wspace=0.20)
    ax_a = fig.add_subplot(gs[0, :])
    ax_b = fig.add_subplot(gs[1, 0])
    ax_c = fig.add_subplot(gs[1, 1])
    style_title(fig, "Policy, coding, cohort, and observation-window timeline")

    dates = pd.to_datetime(policy["date"].map(lambda x: x + "-01" if len(str(x)) == 7 else x))
    labels = [
        "2019-07-04\nCONITEC recommendation",
        "2019-07-24\nSCTIE inclusion decision",
        "2020-12-22\nGM/MS coding and financing",
        "2021-01\nSIGTAP observed window",
    ]
    ys = [0.19, -0.19, 0.19, -0.19]
    ax_a.axhline(0, color=PALETTE["grey"], lw=0.8)
    for d, y, label in zip(dates, ys, labels):
        ax_a.vlines(d, 0, y, color=PALETTE["grey"], lw=0.7)
        ax_a.scatter(d, 0, s=42, color=PALETTE["blue"], zorder=3)
        ax_a.text(d, y + (0.02 if y > 0 else -0.02), label, ha="center", va="bottom" if y > 0 else "top", fontsize=6.5, linespacing=1.05)
    ax_a.set_ylim(-0.40, 0.40)
    ax_a.set_xlim(pd.Timestamp("2019-06-01"), pd.Timestamp("2021-03-01"))
    ax_a.set_yticks([])
    ax_a.set_xticks([pd.Timestamp("2020-01-01"), pd.Timestamp("2021-01-01")], ["2020", "2021"])
    ax_a.set_title("Official policy and coding evidence", loc="left", fontsize=8.5, fontweight="bold", pad=7)
    ax_a.text(0, -0.18, "Dates and milestones are transcribed from the hash-locked policy evidence table.", transform=ax_a.transAxes, fontsize=5.6, color=PALETTE["grey"], ha="left")
    ax_a.spines["left"].set_visible(False)
    clean_axis(ax_a)
    panel_letter(ax_a, "a")

    wanted = ["cohort.a_unique_aih", "cohort.b_final_unique_aih", "uptake.left_censored_202101"]
    sub = funnel.set_index("result_id").loc[wanted].reset_index()
    values = sub["value"].astype(float).to_numpy()
    labels_b = ["Cohort A unique AIHs", "Cohort B final unique AIHs", "Prevalent users at the\nJanuary 2021 boundary"]
    colors = [PALETTE["blue"], PALETTE["teal"], PALETTE["orange"]]
    ypos = np.arange(3)[::-1]
    ax_b.barh(ypos, values, color=colors, height=0.64)
    for y, value in zip(ypos, values):
        ax_b.text(value + max(values) * 0.012, y, f"{int(value):,}", va="center", fontsize=6.5)
    ax_b.set_yticks(ypos, labels_b)
    ax_b.set_xlim(0, max(values) * 1.20)
    ax_b.set_xlabel("Records or hospitals")
    ax_b.set_title("Frozen cohort and boundary counts", loc="left", fontsize=8.5, fontweight="bold", pad=7)
    clean_axis(ax_b)
    panel_letter(ax_b, "b")

    ax_c.set_axis_off()
    panel_letter(ax_c, "c")
    ax_c.text(0.02, 1.01, "Observation boundary", transform=ax_c.transAxes, fontsize=8.5, fontweight="bold", color=PALETTE["orange"], va="bottom")
    y = 0.71
    cards = [
        ("JAN 2021", "Observed-window start", PALETTE["blue_soft"], PALETTE["blue"]),
        ("LEFT-CENSORED", "23 prevalent users", PALETTE["orange_soft"], PALETTE["orange"]),
        ("REPORT AS", "First observed coded use", PALETTE["teal_soft"], PALETTE["teal"]),
    ]
    for i, (kicker, body, face, edge) in enumerate(cards):
        rounded_card(ax_c, (0.04, y), 0.84, 0.19, face=face, edge=edge)
        ax_c.add_patch(Circle((0.105, y + 0.095), 0.035, transform=ax_c.transAxes, facecolor=edge, edgecolor="none"))
        ax_c.text(0.105, y + 0.095, str(i + 1), transform=ax_c.transAxes, ha="center", va="center", fontsize=6.2, fontweight="bold", color="white")
        ax_c.text(0.17, y + 0.125, kicker, transform=ax_c.transAxes, fontsize=5.7, fontweight="bold", color=edge, va="center")
        ax_c.text(0.17, y + 0.062, body, transform=ax_c.transAxes, fontsize=6.6, color=PALETTE["dark"], va="center")
        if i < 2:
            ax_c.annotate("", xy=(0.46, y - 0.055), xytext=(0.46, y - 0.005), xycoords=ax_c.transAxes, arrowprops=dict(arrowstyle="-|>", color=PALETTE["grey"], lw=0.8))
        y -= 0.27
    ax_c.text(0.04, 0.01, "Not a national implementation date", transform=ax_c.transAxes, fontsize=6.4, fontweight="bold", color=PALETTE["red"])
    fig.subplots_adjust(left=0.08, right=0.985, top=0.91, bottom=0.10)
    save_figure(fig, "Figure_1", 183, 120)


def build_figure_3() -> None:
    equity = read_csv("figure_3_equity_source_data.csv")
    travel = read_csv("figure_3_travel_source_data.csv")
    fig = plt.figure(constrained_layout=False)
    gs = fig.add_gridspec(2, 2, height_ratios=[1.35, 1.0], width_ratios=[1.1, 1.0], hspace=0.38, wspace=0.24)
    ax_a = fig.add_subplot(gs[0, :])
    ax_b = fig.add_subplot(gs[1, 0])
    ax_c = fig.add_subplot(gs[1, 1])
    style_title(fig, "Equity gradients and realised travel burden")

    for exposure, color, label in [
        ("IVS_2010_population_weighted_quintile", PALETTE["blue"], "IVS (ecological)"),
        ("ANS_supplementary_coverage_population_weighted_quintile", PALETTE["orange"], "Supplementary insurance (ecological)"),
    ]:
        d = equity[equity["exposure"] == exposure].copy()
        x = d["quintile"].astype(float).to_numpy()
        y = d["age_sex_standardised_rate_per_100k"].astype(float).to_numpy()
        lo = d["standardised_lo95"].astype(float).to_numpy()
        hi = d["standardised_hi95"].astype(float).to_numpy()
        ax_a.errorbar(x, y, yerr=[y - lo, hi - y], color=color, marker="o", ms=4.2, lw=1.2, elinewidth=0.8, capsize=4, label=label)
    ax_a.set_xlabel("Contextual quintile")
    ax_a.set_ylabel("AIHs per 100,000 adults")
    ax_a.set_xticks([1, 2, 3, 4, 5])
    ax_a.set_ylim(2.5, 6.25)
    ax_a.set_title("Directly standardised treated-use rates", loc="left", fontsize=8.5, fontweight="bold", pad=15)
    ax_a.text(0, 1.03, "Municipality-year ecological strata; bars show frozen 95% intervals", transform=ax_a.transAxes, fontsize=6.4, va="bottom")
    ax_a.legend(loc="center left", bbox_to_anchor=(1.02, 0.5), fontsize=6.4)
    clean_axis(ax_a)
    panel_letter(ax_a, "a")

    travel["known"] = pd.to_numeric(travel["weighted_n_aih_display"], errors="coerce")
    suppressed = int(travel["known"].isna().sum())
    agg = travel.dropna(subset=["known"]).groupby("travel_bin_label", as_index=False)["known"].sum()
    order = ["NOT_ROUTED", "0-<15", "15-<30", "30-<45", "45-<60", "60-<75", "75-<90", "90-<105", "105-<120", "120-<135", "135-<150", "150-<165", "165-<180", "180+"]
    agg["travel_bin_label"] = pd.Categorical(agg["travel_bin_label"], order, ordered=True)
    agg = agg.sort_values("travel_bin_label")
    x = np.arange(len(agg))
    ax_b.bar(x, agg["known"].astype(float), color=PALETTE["teal"], width=0.86)
    ax_b.axvline(8.5, color=PALETTE["red"], lw=0.9, ls="--")
    ax_b.text(8.58, ax_b.get_ylim()[1] * 0.92, "120 min", color=PALETTE["red"], fontsize=5.6, va="top")
    ax_b.set_xticks(x, [str(v) for v in agg["travel_bin_label"]], rotation=45, ha="right")
    ax_b.set_ylabel("Known weighted AIHs")
    ax_b.set_xlabel("15-minute bin")
    ax_b.set_title("Realised flow-weighted road-time distribution", loc="left", fontsize=8.2, fontweight="bold", pad=15)
    ax_b.text(0, 1.03, f"{suppressed} privacy-suppressed strata are not imputed", transform=ax_b.transAxes, fontsize=6.2, va="bottom")
    clean_axis(ax_b)
    panel_letter(ax_b, "b")

    ax_c.set_axis_off()
    panel_letter(ax_c, "c")
    ax_c.text(0.02, 1.01, "Interpretation boundary", transform=ax_c.transAxes, fontsize=8.5, fontweight="bold", color=PALETTE["orange"], va="bottom")
    cards = [
        ("ECOLOGICAL CONTEXT", "Population-level gradients", "not individual effects", PALETTE["blue_soft"], PALETTE["blue"]),
        ("OBSERVED TREATED FLOWS", "Municipal-anchor road time", "not population access", PALETTE["teal_soft"], PALETTE["teal"]),
    ]
    y = 0.61
    for kicker, line1, line2, face, edge in cards:
        rounded_card(ax_c, (0.04, y), 0.88, 0.27, face=face, edge=edge)
        ax_c.add_patch(Circle((0.13, y + 0.135), 0.048, transform=ax_c.transAxes, facecolor=edge, edgecolor="none"))
        ax_c.text(0.13, y + 0.135, "●", transform=ax_c.transAxes, color="white", fontsize=7, ha="center", va="center")
        ax_c.text(0.21, y + 0.19, kicker, transform=ax_c.transAxes, fontsize=5.8, fontweight="bold", color=edge, va="center")
        ax_c.text(0.21, y + 0.12, line1, transform=ax_c.transAxes, fontsize=6.6, color=PALETTE["dark"], va="center")
        ax_c.text(0.21, y + 0.055, line2, transform=ax_c.transAxes, fontsize=6.2, fontweight="bold", color=PALETTE["red"], va="center")
        y -= 0.34
    rounded_card(ax_c, (0.04, 0.00), 0.88, 0.16, face=PALETTE["grey_soft"], edge=PALETTE["line"])
    ax_c.text(0.48, 0.08, "Descriptive / associational evidence only", transform=ax_c.transAxes, ha="center", va="center", fontsize=6.4, fontweight="bold", color=PALETTE["dark"])
    fig.subplots_adjust(left=0.08, right=0.985, top=0.91, bottom=0.11)
    save_figure(fig, "Figure_3", 183, 120)


def build_figure_4() -> None:
    coverage = read_csv("figure_4_national_regional_coverage_source_data.csv")
    municipality = read_csv("figure_4_municipality_coverage_source_data.csv")
    national = coverage[(coverage["scope"] == "national") & (coverage["status"] == "EVALUATED")].copy()
    fig = plt.figure(constrained_layout=False)
    gs = fig.add_gridspec(2, 2, height_ratios=[1.4, 1.0], width_ratios=[1.0, 1.0], hspace=0.38, wspace=0.24)
    ax_a = fig.add_subplot(gs[0, :])
    ax_b = fig.add_subplot(gs[1, 0])
    ax_c = fig.add_subplot(gs[1, 1])
    style_title(fig, "Potential population access within 120 and 180 minutes")

    years = national["year"].astype(float).astype(int).to_numpy()
    s120 = national["adult_population_share_120"].astype(float).to_numpy()
    s180 = national["adult_population_share_180"].astype(float).to_numpy()
    ax_a.plot(years, s120, color=PALETTE["blue"], marker="o", ms=4.2, lw=1.4, label="120 minutes")
    ax_a.plot(years, s180, color=PALETTE["teal"], marker="o", ms=4.2, lw=1.4, label="Cumulative 180 minutes")
    for x, y in [(years[-1], s120[-1]), (years[-1], s180[-1])]:
        ax_a.text(x + 0.05, y, f"{100*y:.1f}%", va="center", fontsize=6.2)
    ax_a.set_xticks(years)
    ax_a.set_ylim(0, 1.05)
    ax_a.set_yticks([0, 0.25, 0.5, 0.75, 1.0], ["0%", "25%", "50%", "75%", "100%"])
    ax_a.set_ylabel("Population share")
    ax_a.set_xlabel("Year")
    ax_a.set_title("Adult-population potential access", loc="left", fontsize=8.5, fontweight="bold", pad=7)
    ax_a.legend(loc="center left", bbox_to_anchor=(1.02, 0.5), fontsize=6.4)
    clean_axis(ax_a)
    panel_letter(ax_a, "a")

    latest = municipality[municipality["year"].astype(float).astype(int) == municipality["year"].astype(float).astype(int).max()].copy()
    anchor = latest["anchor_available"].astype(str).str.lower().eq("true")
    has120 = latest["has_provider_120"].astype(str).str.lower().eq("true")
    has180 = latest["has_provider_180"].astype(str).str.lower().eq("true")
    latest["category"] = np.select(
        [~anchor, has120, has180],
        ["Missing anchor", "≤120 minutes", "120–180 minutes"],
        default=">180 minutes",
    )
    latest["adult_population"] = latest["adult_population"].astype(float)
    shares = latest.groupby("category")["adult_population"].sum() / latest["adult_population"].sum()
    categories = ["≤120 minutes", ">180 minutes", "120–180 minutes", "Missing anchor"]
    vals = [float(shares.get(c, 0)) for c in categories]
    colors = [PALETTE["blue"], PALETTE["orange"], PALETTE["teal"], PALETTE["grey"]]
    y = np.arange(len(categories))[::-1]
    ax_b.barh(y, vals, color=colors, height=0.68)
    for yy, val in zip(y, vals):
        label = "<0.1%" if 0 < val < 0.0005 else f"{100*val:.1f}%"
        ax_b.text(val + 0.006, yy, label, va="center", fontsize=6.4)
    ax_b.set_yticks(y, categories)
    ax_b.set_xlim(0, max(vals) * 1.19)
    ticks = [0, 0.25, 0.5, 0.75]
    ax_b.set_xticks(ticks, [f"{int(100*t)}%" for t in ticks])
    ax_b.set_xlabel("Adult-population share")
    ax_b.set_title("2025 anchored population categories", loc="left", fontsize=8.2, fontweight="bold", pad=7)
    clean_axis(ax_b)
    panel_letter(ax_b, "b")

    ax_c.set_axis_off()
    panel_letter(ax_c, "c")
    ax_c.text(0.02, 1.01, "Frozen-input boundary", transform=ax_c.transAxes, fontsize=8.5, fontweight="bold", color=PALETTE["orange"], va="bottom")
    chip(ax_c, 0.04, 0.79, "EVALUATED", face=PALETTE["teal_soft"], color=PALETTE["teal"], width=0.25)
    ax_c.text(0.33, 0.837, "National 120- and 180-minute trends", transform=ax_c.transAxes, va="center", fontsize=6.4, color=PALETTE["dark"])
    chip(ax_c, 0.04, 0.63, "NOT EVALUATED", face=PALETTE["red_soft"], color=PALETTE["red"], width=0.31)
    missing = ["Geographic maps", "Regional coverage", "Vulnerability gap"]
    for i, label in enumerate(missing):
        y0 = 0.58 - i * 0.12
        ax_c.add_patch(Circle((0.095, y0), 0.022, transform=ax_c.transAxes, facecolor=PALETTE["red"], edgecolor="none"))
        ax_c.text(0.095, y0, "×", transform=ax_c.transAxes, ha="center", va="center", fontsize=6.2, fontweight="bold", color="white")
        ax_c.text(0.14, y0, label, transform=ax_c.transAxes, fontsize=6.3, va="center", color=PALETTE["dark"])
    rounded_card(ax_c, (0.04, 0.04), 0.88, 0.22, face=PALETTE["orange_soft"], edge=PALETTE["orange"])
    ax_c.text(0.08, 0.205, "Cumulative 180-minute rule retained", transform=ax_c.transAxes, fontsize=6.4, fontweight="bold", color=PALETTE["orange"], va="center")
    ax_c.text(0.08, 0.125, "Potential access ≠ realised use, capacity,", transform=ax_c.transAxes, fontsize=6.1, color=PALETTE["dark"], va="center")
    ax_c.text(0.08, 0.073, "quality, or need", transform=ax_c.transAxes, fontsize=6.1, color=PALETTE["dark"], va="center")
    fig.subplots_adjust(left=0.08, right=0.985, top=0.91, bottom=0.10)
    save_figure(fig, "Figure_4", 183, 135)


def metric_formatter(metric: str, value: float) -> str:
    if metric in {"risk_p10", "risk_p90"}:
        return f"{100*value:.3f}%"
    if metric == "rd":
        return f"{100*value:.4f} pp"
    return f"{value:.3f}"


def build_figure_6() -> None:
    points = read_csv("figure_6_adjusted_point_estimates_source_data.csv")
    boot = read_csv("figure_6_bootstrap_validity_source_data.csv")
    fig = plt.figure(constrained_layout=False)
    outer = fig.add_gridspec(2, 2, height_ratios=[1.25, 1.0], width_ratios=[1.0, 1.0], hspace=0.38, wspace=0.20)
    top = outer[0, :].subgridspec(1, 4, wspace=0.32)
    axes = [fig.add_subplot(top[0, i]) for i in range(4)]
    ax_b = fig.add_subplot(outer[1, 0])
    ax_c = fig.add_subplot(outer[1, 1])
    style_title(fig, "Hospital treatment volume and in-hospital death", "Associational point estimates; no causal interpretation.")
    metric_titles = [
        ("risk_p10", "Death risk at volume p10"),
        ("risk_p90", "Death risk at volume p90"),
        ("rd", "Risk difference (p90 − p10)"),
        ("rr", "Risk ratio (p90 / p10)"),
    ]
    estimators = ["AS_mean", "MPL_Jeffreys"]
    colors = [PALETTE["blue"], PALETTE["teal"]]
    for idx, (ax, (metric, title)) in enumerate(zip(axes, metric_titles)):
        vals = []
        for estimator in estimators:
            vals.append(float(points[(points["metric"] == metric) & (points["estimator"] == estimator)]["value"].iloc[0]))
        ax.scatter([0, 1], vals, s=38, color=colors, zorder=3)
        for x, value in zip([0, 1], vals):
            ax.text(x, value, metric_formatter(metric, value), ha="center", va="bottom", fontsize=5.6, color=PALETTE["dark"])
        if metric == "rd":
            ax.axhline(0, color=PALETTE["grey"], lw=0.7, ls="--")
            ax.set_ylim(-0.00045, 0.00035)
        elif metric == "rr":
            ax.axhline(1, color=PALETTE["grey"], lw=0.7, ls="--")
            ax.set_ylim(0.95, 1.03)
        else:
            ax.set_ylim(0.0139, 0.01475)
        ax.set_xticks([])
        ax.set_title(title, fontsize=6.5, fontweight="bold", pad=7, bbox=dict(facecolor=PALETTE["grey_soft"], edgecolor=PALETTE["line"], boxstyle="round,pad=0.35"))
        if idx == 0:
            ax.set_ylabel("Point estimate")
        clean_axis(ax)
    panel_letter(axes[0], "a")
    fig.text(0.08, 0.902, "Associational point estimates only", fontsize=8.2, fontweight="bold", va="bottom")
    fig.text(0.08, 0.875, "Formal intervals: NOT_EVALUATED", fontsize=6.2, va="bottom")
    fig.legend(
        [
            Line2D([0], [0], marker="o", color="none", markerfacecolor=PALETTE["blue"], markeredgecolor=PALETTE["blue"], markersize=5),
            Line2D([0], [0], marker="o", color="none", markerfacecolor=PALETTE["teal"], markeredgecolor=PALETTE["teal"], markersize=5),
        ],
        estimators,
        loc="upper right",
        bbox_to_anchor=(0.975, 0.955),
        ncol=2,
        fontsize=6.0,
        handletextpad=0.4,
        columnspacing=1.0,
    )

    valid = boot[boot["result_id"].str.contains("valid_replicates")].copy()
    valid_vals = valid["value"].astype(float).to_numpy()
    ax_b.bar([0, 1], valid_vals, color=colors, width=0.62)
    ax_b.axhline(2000, color=PALETTE["grey"], lw=0.9, ls="--")
    for x, value in zip([0, 1], valid_vals):
        ax_b.text(x, value + 55, f"{int(value)}/2000", ha="center", va="bottom", fontsize=6.5)
    ax_b.set_ylim(0, 2150)
    ax_b.set_xticks([0, 1], estimators)
    ax_b.set_ylabel("Valid replicates")
    ax_b.set_title("Bootstrap valid-replicate yield", loc="left", fontsize=8.2, fontweight="bold", pad=15)
    ax_b.text(0, 1.03, "Prespecified gate: DOWNGRADE", transform=ax_b.transAxes, fontsize=6.2, va="bottom", color=PALETTE["red"])
    clean_axis(ax_b)
    panel_letter(ax_b, "b")

    ax_c.set_axis_off()
    panel_letter(ax_c, "c")
    ax_c.text(0.02, 1.01, "Inference gate", transform=ax_c.transAxes, fontsize=8.5, fontweight="bold", color=PALETTE["red"], va="bottom")
    steps = [
        ("2,000", "planned replicates", PALETTE["grey_soft"], PALETTE["grey"]),
        ("618 / 617", "valid replicates", PALETTE["orange_soft"], PALETTE["orange"]),
        ("DOWNGRADE", "bootstrap gate", PALETTE["red_soft"], PALETTE["red"]),
        ("NOT_EVALUATED", "formal intervals", PALETTE["red_soft"], PALETTE["red"]),
    ]
    xs = [0.02, 0.265, 0.51, 0.755]
    for i, (x, (top_text, bottom_text, face, edge)) in enumerate(zip(xs, steps)):
        rounded_card(ax_c, (x, 0.57), 0.205, 0.25, face=face, edge=edge)
        ax_c.text(x + 0.1025, 0.72, top_text, transform=ax_c.transAxes, ha="center", va="center", fontsize=6.3 if i != 3 else 4.8, fontweight="bold", color=edge)
        ax_c.text(x + 0.1025, 0.62, bottom_text, transform=ax_c.transAxes, ha="center", va="center", fontsize=5.5, color=PALETTE["dark"])
        if i < 3:
            ax_c.annotate("", xy=(xs[i + 1] - 0.01, 0.695), xytext=(x + 0.218, 0.695), xycoords=ax_c.transAxes, arrowprops=dict(arrowstyle="-|>", color=PALETTE["grey"], lw=0.75))
    rounded_card(ax_c, (0.04, 0.16), 0.88, 0.25, face=PALETTE["blue_soft"], edge=PALETTE["blue"])
    chip(ax_c, 0.075, 0.29, "RETAINED", face=PALETTE["teal_soft"], color=PALETTE["teal"], width=0.22)
    ax_c.text(0.33, 0.338, "Associational point estimates only", transform=ax_c.transAxes, fontsize=6.5, fontweight="bold", color=PALETTE["dark"], va="center")
    ax_c.text(0.075, 0.225, "No rerun or model relaxation; no causal or equivalence claim", transform=ax_c.transAxes, fontsize=6.1, color=PALETTE["dark"], va="center")
    fig.subplots_adjust(left=0.08, right=0.985, top=0.83, bottom=0.10)
    save_figure(fig, "Figure_6", 183, 120)


def copy_unchanged_figures() -> None:
    mapping = {"Figure_2": "Figure_2", "Figure_5": "Figure_5", "Supplementary_Figure_1": "Supplementary_Figure_1"}
    for source_stem, target_stem in mapping.items():
        for ext in ["svg", "pdf", "tiff"]:
            src = SOURCE_RENDERED / f"{source_stem}.{ext}"
            shutil.copy2(src, FIGURES / f"{target_stem}.{ext}")
        shutil.copy2(SOURCE_PREVIEWS / f"{source_stem}.png", FIGURES / f"{target_stem}.png")


FIGURE_LEGENDS = """
## Figure legends

**Figure 1. Policy, coding, cohort, and observation-window timeline.** (a) Hash-locked policy and coding milestones. (b) Frozen counts for Cohort A, Cohort B, and hospitals active at the January 2021 observation boundary. (c) Visual boundary map distinguishing the observed-window start, left-censored prevalent users, and first observed coded use. The January 2021 boundary is not a national implementation date.

**Figure 2. National observed uptake, diffusion, and service maintenance.** Monthly and annual descriptive therapeutic ERCP use, first observed coded use by year, and billing-defined maintenance among hospitals with an evaluable completed 12-month window. First observed coded use is not a true adoption date, and coded maintenance is not audited service capacity.

**Figure 3. Equity gradients and realised travel burden.** (a) Directly standardised treated-use rates across ecological IVS and supplementary-insurance quintiles with frozen 95% intervals. (b) Flow-weighted road-time distribution among observed treated flows; privacy-suppressed strata are not imputed. (c) Interpretation map separating ecological context and observed routed travel from individual effects and population access.

**Figure 4. Potential population access within 120 and 180 minutes.** (a) National adult-population shares within 120 and cumulative 180 road-minutes of observed provider municipalities. (b) Population categories for 2025 using anchored municipalities. (c) Frozen-input status map. Geographic maps, regional coverage, and the vulnerability gap remain NOT_EVALUATED; potential access is not realised use, capacity, quality, or need.

**Figure 5. Intermunicipal patient-flow network summaries and referral concentration.** Non-spatial summaries of the privacy-suppressed patient-flow network, service-area modal-destination shares, and centrality distributions. Projection edges are not patient-to-patient links, and cross-municipal flow does not establish poor access or inappropriate referral.

**Figure 6. Hospital treatment volume and in-hospital death.** (a) Standardised risks and contrasts at hospital-volume p10 and p90 from the two prespecified bias-reduced estimators. (b) Valid bootstrap-replicate yield. (c) Prespecified inference-gate cascade. Point estimates are associational only; the bootstrap gate is DOWNGRADE and formal intervals are NOT_EVALUATED.
""".strip()


SUPP_FIGURE_LEGEND = (
    "Supplementary Figure 1. Network resilience under targeted hub removal versus random removal. "
    "Targeted removal of prespecified high-in-strength hubs is compared with seeded random-removal "
    "benchmarks at the 120- and 180-minute thresholds. Random envelopes are simulation benchmarks, "
    "not confidence intervals. This is a structural stress test, not a real-world intervention counterfactual."
)


def transform_main_md() -> str:
    text = SOURCE_MD.read_text(encoding="utf-8")
    text = text.replace("Supplementary Table S1", "Supplementary Table 2")
    text = text.replace("Supplementary Table S2", "Supplementary Table 3")
    text = text.replace("Supplementary Table S3", "Supplementary Table 4")
    text = text.replace(
        "Supplementary Sections S3, S4, and S9",
        "the supplementary sections on cohort construction, missing-data handling, and the supportive outcome analysis",
    )
    text = re.sub(r"(?<!Supplementary )Figure 7", "Supplementary Figure 1", text)
    text = re.sub(r"(?<!Supplementary )Table 3", "Supplementary Table 5", text)
    text = re.sub(r"(?<!Supplementary )Table 4", "Supplementary Table 6", text)
    text = text.replace("Supplementary Section S9", "the supplementary section ‘Aim 4 supportive outcome analysis and bootstrap failure taxonomy’")
    text = text.replace("## References", FIGURE_LEGENDS + "\n\n## References")
    return text


SECTION_NAME_MAP = {
    "S1": "Policy, coding, and protocol chronology",
    "S2": "Data sources and acquisition boundary",
    "S3": "Code lists, linkage, and cohort construction",
    "S4": "Variable availability and missing-data handling",
    "S5": "Cross-aim estimand, denominator, missingness, and gate map",
    "S6": "Aim 1 definitions and diagnostics",
    "S7": "Aim 2 ecological equity, routing, and potential reach",
    "S8": "Aim 3 patient-flow network, service areas, privacy, and resilience",
    "S9": "Aim 4 supportive outcome analysis and bootstrap failure taxonomy",
    "S10": "Stage 6 quasi-causal gate",
    "S11": "Figure, table, and numeric traceability",
    "S12": "Display-rounding policy",
    "S13": "RECORD and STROBE locator checklist",
    "S14": "Reproducibility, availability, and declarations",
    "S15": "Source and boundary ledger",
}


def replace_section_locators(text: str) -> str:
    replacements = {
        "Supplement S2, S6–S9": "the supplementary sections on data sources and the aim-specific methods",
        "Supplement S2 and S8": "the supplementary sections ‘Data sources and acquisition boundary’ and ‘Aim 3 patient-flow network, service areas, privacy, and resilience’",
        "Supplement S3–S4": "the supplementary sections ‘Code lists, linkage, and cohort construction’ and ‘Variable availability and missing-data handling’",
        "Supplement S6–S10": "the supplementary aim-specific and quasi-causal-gate sections",
        "Supplement S6–S9": "the supplementary aim-specific sections",
        "Supplement S5–S9": "the supplementary estimand map and aim-specific sections",
        "Supplement S5 and S9": "the supplementary estimand map and Aim 4 section",
        "Supplement S5": "the supplementary estimand map",
        "Supplement S14": "the supplementary reproducibility and declarations section",
        "Supplement S3": "the supplementary cohort-construction section",
        "Supplement S4": "the supplementary missing-data section",
        "Supplement S2": "the supplementary data-sources section",
        "Supplement S8": "the supplementary patient-flow network section",
        "Supplement S9": "the supplementary Aim 4 section",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace(
        "Results; Figures 3–7; Tables 3–4",
        "Results; Figures 3–6, Supplementary Figure 1, and Supplementary Tables 5–6",
    )
    return text


def strip_markdown_tables_and_renumber(text: str) -> tuple[str, list[dict]]:
    lines = text.splitlines()
    captured: list[dict] = []
    output: list[str] = []
    i = 0
    while i < len(lines):
        match = re.match(r"^### Supplementary Table S?([1-3])(?::\s*(.*))?$", lines[i])
        if match:
            number = int(match.group(1)) + 1
            title = match.group(2) or SECTION_NAME_MAP["S5"]
            output.append(f"### Supplementary Table {number}: {title}")
            output.append("")
            output.append(f"The editable table is supplied in `table.docx` as Supplementary Table {number}.")
            output.append("")
            i += 1
            while i < len(lines) and not lines[i].startswith("|"):
                if lines[i].strip():
                    break
                i += 1
            table_lines: list[str] = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            captured.append({"number": number, "title": title, "lines": table_lines})
            continue
        if lines[i].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            number = 1
            title = "Data sources and acquisition boundaries"
            output.append(f"### Supplementary Table {number}: {title}")
            output.append("")
            output.append(f"The editable table is supplied in `table.docx` as Supplementary Table {number}.")
            output.append("")
            captured.append({"number": number, "title": title, "lines": table_lines})
            continue
        output.append(lines[i])
        i += 1
    return "\n".join(output), captured


def transform_supplement_md() -> tuple[str, list[dict]]:
    text = SOURCE_SUPP_MD.read_text(encoding="utf-8")
    text = re.sub(r"^## S\d+\.\s+", "## ", text, flags=re.MULTILINE)
    text = text.replace("Supplementary Table S1", "Supplementary Table 1")
    text = text.replace("Supplementary Table S2", "Supplementary Table 2")
    text = text.replace("Supplementary Table S3", "Supplementary Table 3")
    text = re.sub(r"(?<!Supplementary )Figure 7", "Supplementary Figure 1", text)
    text = re.sub(r"(?<!Supplementary )Table 3", "Supplementary Table 5", text)
    text = re.sub(r"(?<!Supplementary )Table 4", "Supplementary Table 6", text)
    text = text.replace("Figures 3–7; Supplementary Tables 5–6", "Figures 3–6, Supplementary Figure 1, and Supplementary Tables 5–6")
    text = replace_section_locators(text)
    text, tables = strip_markdown_tables_and_renumber(text)
    transfer_note = """
## Supplementary displays transferred from the main manuscript

To keep the main manuscript at eight displays, the structural-resilience figure and the two detailed result tables were transferred without changing their values, labels, denominators, evidence grades, or scientific boundaries. They are now Supplementary Figure 1 and Supplementary Tables 5–6. All six supplementary tables are supplied as editable tables in `table.docx`.
""".strip()
    marker = "## Display-rounding policy"
    text = text.replace(marker, transfer_note + "\n\n" + marker)
    return text, tables


def set_font(run, name: str, size: float, *, bold: bool | None = None, italic: bool | None = None, color: str = "000000") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(.*?)\*(?!\*)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.strip()


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, "Times New Roman", 9, color="666666")
    for kind, value in [("begin", None), (None, " PAGE "), ("separate", None), (None, "1"), ("end", None)]:
        if kind:
            element = OxmlElement("w:fldChar")
            element.set(qn("w:fldCharType"), kind)
        elif value == " PAGE ":
            element = OxmlElement("w:instrText")
            element.set(qn("xml:space"), "preserve")
            element.text = value
        else:
            element = OxmlElement("w:t")
            element.text = value
        run._r.append(element)


def configure_document(doc: Document, *, kind: str, landscape: bool = False) -> None:
    body_size = 12 if kind == "main" else 10.5
    spacing = 2.0 if kind == "main" else 1.15
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_after = Pt(6 if kind == "main" else 4)
    normal.paragraph_format.line_spacing = spacing
    for style_name, size, before, after in [("Title", 16, 0, 12), ("Heading 1", 14, 16, 8), ("Heading 2", 12.5, 12, 6), ("Heading 3", 11.5, 10, 4)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0
    for section in doc.sections:
        section.page_width = Inches(11 if landscape else 8.5)
        section.page_height = Inches(8.5 if landscape else 11)
        section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
        margin = 0.6 if landscape else (1.0 if kind == "main" else 0.7)
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)
        section.header_distance = Inches(0.45)
        section.footer_distance = Inches(0.45)
        header = section.header.paragraphs[0]
        header.text = "Blinded manuscript — DATASUS ERCP" if kind == "main" else ("Tables — DATASUS ERCP" if landscape else "Supplementary material — DATASUS ERCP")
        for run in header.runs:
            set_font(run, "Times New Roman", 8.5, italic=True, color="666666")
        add_page_field(section.footer.paragraphs[0])


def add_markdown(doc: Document, text: str, *, kind: str) -> None:
    lines = text.splitlines()
    buffer: list[str] = []
    in_refs = False

    def flush() -> None:
        nonlocal buffer
        if not buffer:
            return
        joined = buffer[0].strip()
        for line in buffer[1:]:
            joined += line.strip() if joined.endswith("-") else " " + line.strip()
        style = "Normal"
        p = doc.add_paragraph(style=style)
        p.paragraph_format.widow_control = True
        r = p.add_run(clean_inline(joined))
        set_font(r, "Times New Roman", 12 if kind == "main" else 10.5)
        if in_refs and re.match(r"^\d+\. ", joined):
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_font(run, "Times New Roman", 9.5)
        buffer = []

    i = 0
    while i < len(lines):
        line = lines[i]
        match = re.match(r"^(#{1,3})\s+(.*)$", line)
        if match:
            flush()
            level = len(match.group(1))
            heading = clean_inline(match.group(2))
            if level == 1 and len(doc.paragraphs) == 0:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(heading)
                set_font(r, "Times New Roman", 16, bold=True)
            else:
                style = "Heading 1" if level == 2 else "Heading 2"
                doc.add_paragraph(heading, style=style)
            in_refs = heading == "References"
        elif re.match(r"^[-*]\s+", line):
            flush()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.45)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            r = p.add_run(clean_inline(re.sub(r"^[-*]\s+", "", line)))
            set_font(r, "Times New Roman", 12 if kind == "main" else 10.5)
        elif not line.strip():
            flush()
        else:
            buffer.append(line)
        i += 1
    flush()


def finalize_properties(doc: Document, title: str) -> None:
    doc.core_properties.title = title
    doc.core_properties.subject = "Journal-neutral blinded editor package"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa: list[int], table_width_dxa: int) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(table_width_dxa))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))


def parse_markdown_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in lines:
        cells = [clean_inline(x) for x in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return [], []
    return rows[0], rows[1:]


def format_value(row: pd.Series) -> str:
    raw = str(row.get("value", ""))
    vtype = str(row.get("value_type", ""))
    unit = str(row.get("unit", ""))
    if vtype == "string" or not re.fullmatch(r"[-+0-9.eE]+", raw):
        return raw.replace("_", " ")
    value = float(raw)
    if unit == "proportion":
        if "anchor_coverage" in str(row.get("result_id", "")):
            return f"{100 * value:.4f}%"
        return f"{100 * value:.1f}%"
    if unit in {"AIHs", "hospitals", "deaths", "replicates", "rows", "nodes", "edges", "municipalities", "records"}:
        return f"{int(round(value)):,} {unit}"
    if unit == "minutes":
        return f"{value:.1f} minutes"
    if unit == "rate ratio":
        return f"{value:.2f}"
    if "100,000" in unit:
        return f"{value:.2f} {unit}"
    return f"{value:,.4g} {unit}".strip()


def format_denominator(row: pd.Series) -> str:
    raw = row.get("denominator_value")
    if pd.isna(raw) or str(raw) in {"NA", "nan", "None"}:
        return "—"
    value = float(raw)
    unit = str(row.get("denominator_unit", ""))
    return f"{int(round(value)):,} {unit}".strip()


def add_word_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int], *, font_size: float = 7.2) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_font(run, "Arial", font_size, bold=True)
    set_repeat_header(table.rows[0])
    set_cant_split(table.rows[0])
    for i, values in enumerate(rows, start=1):
        row = table.add_row()
        set_cant_split(row)
        for j, value in enumerate(values):
            cell = row.cells[j]
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i % 2 == 0:
                set_cell_shading(cell, "F7F9FA")
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_font(run, "Arial", font_size)
    set_table_geometry(table, widths_dxa, sum(widths_dxa))


MAIN_TABLE_TITLES = {
    1: "Cohort, hospital, municipality, and regional characteristics",
    2: "Observed uptake and maintenance of therapeutic ERCP",
}
SUPP_TABLE_TITLES = {
    5: "Equity, travel time, potential access, and patient-flow concentration",
    6: "Adjusted hospital attributes and in-hospital outcome associations",
}


def add_registry_table(doc: Document, source_number: int, display_label: str, title: str) -> None:
    doc.add_paragraph(f"{display_label}. {title}", style="Heading 1")
    data = pd.read_csv(SOURCE_RENDERED / f"Table_{source_number}.csv", dtype=str, keep_default_na=True)
    rows = []
    for _, row in data.iterrows():
        rows.append(
            [
                str(row["display_label"]),
                format_value(row),
                format_denominator(row),
                str(row["evidence_level"]).replace("_", " "),
                str(row["limitation"]),
            ]
        )
    add_word_table(doc, ["Measure", "Value", "Denominator", "Evidence", "Key limitation"], rows, [2800, 1700, 2100, 1450, 5150], font_size=7.0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Values are frozen registry outputs; wording and display precision do not alter the accepted estimands or gates.")
    set_font(r, "Arial", 8, italic=True, color="666666")


def build_table_docx(captured_tables: list[dict]) -> None:
    doc = Document()
    configure_document(doc, kind="tables", landscape=True)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tables for editor review")
    set_font(r, "Times New Roman", 16, bold=True)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = note.add_run("Main manuscript: Tables 1–2. Supplementary material: Supplementary Tables 1–6.")
    set_font(rr, "Times New Roman", 10.5, italic=True, color="555555")
    doc.add_page_break()
    for i in [1, 2]:
        add_registry_table(doc, i, f"Table {i}", MAIN_TABLE_TITLES[i])
        doc.add_page_break()
    for item in sorted(captured_tables, key=lambda x: x["number"]):
        headers, rows = parse_markdown_table(item["lines"])
        doc.add_paragraph(f"Supplementary Table {item['number']}. {item['title']}", style="Heading 1")
        if len(headers) == 7:
            widths = [1550, 1550, 1450, 2100, 1250, 3150, 1200]
            size = 6.5
        elif len(headers) == 3:
            widths = [2550, 6600, 1900]
            size = 7.3
        else:
            total = 11050
            widths = [total // len(headers)] * len(headers)
            widths[-1] += total - sum(widths)
            size = 7.0
        add_word_table(doc, headers, rows, widths, font_size=size)
        doc.add_page_break()
    add_registry_table(doc, 3, "Supplementary Table 5", SUPP_TABLE_TITLES[5])
    doc.add_page_break()
    add_registry_table(doc, 4, "Supplementary Table 6", SUPP_TABLE_TITLES[6])
    finalize_properties(doc, "Tables for editor review — DATASUS ERCP")
    doc.save(TABLE_DOCX)


def build_manuscript_docx(main_text: str) -> None:
    doc = Document()
    configure_document(doc, kind="main")
    add_markdown(doc, main_text, kind="main")
    finalize_properties(doc, "Observed uptake, geographic reach, and patient-flow structure of therapeutic ERCP in Brazil")
    doc.save(MAIN_DOCX)


def build_supplement_docx(supp_text: str) -> None:
    doc = Document()
    configure_document(doc, kind="supplement")
    add_markdown(doc, supp_text, kind="supplement")
    doc.add_page_break()
    doc.add_paragraph("Supplementary Figure 1", style="Heading 1")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGURES / "Supplementary_Figure_1.png"), width=Inches(6.65))
    cap = doc.add_paragraph()
    cap.paragraph_format.keep_with_next = False
    rr = cap.add_run(SUPP_FIGURE_LEGEND)
    set_font(rr, "Times New Roman", 9.5)
    finalize_properties(doc, "Supplementary material — DATASUS ERCP nationwide administrative study")
    doc.save(SUPP_DOCX)


def write_figure_contract() -> None:
    text = """# Figure redesign contract

Core conclusion: The four boundary panels communicate limits as structured evidence maps rather than detached prose blocks, without adding or altering any result.

Figure archetype: Integrated journal-width composite; redesigned panel c uses status chips, linked cards, and compact evidence-gate diagrams.

Target output: Journal-neutral editor package; SVG, PDF, 600-dpi TIFF, and PNG preview.

Backend: Python (matplotlib) only.

Panel map:
- Figure 1c: observation boundary → left-censoring → first observed coded use, with explicit non-implementation boundary.
- Figure 3c: ecological context and observed treated-flow road time mapped to their prohibited individual/access interpretations.
- Figure 4c: evaluated national trend separated from NOT_EVALUATED geographic, regional, and vulnerability components.
- Figure 6c: planned bootstrap → valid yield → DOWNGRADE → formal intervals NOT_EVALUATED, with retained associational point estimates.

Evidence hierarchy: Existing quantitative panels remain primary; boundary cards are interpretive safeguards.

Statistics: No refitting, recomputation, interpolation, or new inference. Frozen source CSVs only.

Image integrity: Original Figures 2, 5, and 7 are copied unchanged. All modified figures are exported from one Python source with editable SVG text.

Reviewer risk: Boundary graphics must not soften DOWNGRADE/NOT_EVALUATED status or convert descriptive/associational evidence into causal claims.
"""
    (TARGET / "figure_redesign_contract.md").write_text(text, encoding="utf-8")


def write_package_note() -> None:
    note = """# Editor package contents

The main manuscript contains exactly eight main displays: Figures 1–6 and Tables 1–2.

The former Figure 7 and Tables 3–4 were transferred without scientific changes:

- former Figure 7 → Supplementary Figure 1;
- former Table 3 → Supplementary Table 5;
- former Table 4 → Supplementary Table 6.

All tables are collected as editable Word tables in `table.docx`. The `Figures` folder contains submission-quality SVG, PDF, 600-dpi TIFF, and PNG preview files. Supplementary section headings use topical names and no S1/S2-style chapter numbers.

The package is blinded because author-owned metadata was not available. No author, ethics, funding, conflict, CRediT, repository DOI, licence, or target-journal fact has been invented.
"""
    (TARGET / "README_editor_package.md").write_text(note, encoding="utf-8")


def write_manifest() -> None:
    files = []
    for path in sorted(TARGET.rglob("*")):
        relative = path.relative_to(TARGET)
        is_qa_intermediate = any(part.startswith("_qa") for part in relative.parts)
        if path.is_file() and not is_qa_intermediate and path.name != "editor_package_manifest_v1.json":
            files.append({"path": path.relative_to(TARGET).as_posix(), "bytes": path.stat().st_size, "sha256": sha256(path)})
    payload = {
        "schema_version": "editor_package_manifest_v1",
        "generated_at": datetime.now(timezone.utc).astimezone().isoformat(),
        "source_release": "final blinded manuscript and frozen publication displays",
        "display_contract": {
            "main_figures": [1, 2, 3, 4, 5, 6],
            "main_tables": [1, 2],
            "main_display_count": 8,
            "supplementary_figure_mapping": {"former Figure 7": "Supplementary Figure 1"},
            "supplementary_table_mapping": {"former Table 3": "Supplementary Table 5", "former Table 4": "Supplementary Table 6"},
            "supplement_section_numbering": "topical headings; no S1/S2 chapter numbering",
        },
        "scientific_guards": {
            "data_raw": "not accessed",
            "model_rerun": False,
            "registry_permissions": "unchanged",
            "claim_level": "descriptive_and_associational",
            "aim2_primary_family": "DOWNGRADE",
            "aim4_bootstrap": "DOWNGRADE",
            "aim4_formal_intervals": "NOT_EVALUATED",
            "stage6_results": "excluded",
        },
        "files": files,
    }
    (TARGET / "editor_package_manifest_v1.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    global TARGET, FIGURES, QA, MAIN_MD, SUPP_MD, MAIN_DOCX, SUPP_DOCX, TABLE_DOCX
    parser = argparse.ArgumentParser(description="Rebuild the enhanced editor figures from frozen aggregate source data.")
    parser.add_argument("--output-dir", default=str(TARGET), help="New output directory; must be absent or empty")
    args = parser.parse_args()
    TARGET = Path(args.output_dir).resolve()
    FIGURES = TARGET / "Figures"
    QA = TARGET / "_qa"
    MAIN_MD = TARGET / "unused_main.md"
    SUPP_MD = TARGET / "unused_supplement.md"
    MAIN_DOCX = TARGET / "unused_main.docx"
    SUPP_DOCX = TARGET / "unused_supplement.docx"
    TABLE_DOCX = TARGET / "unused_tables.docx"
    ensure_empty_target()
    build_figure_1()
    build_figure_3()
    build_figure_4()
    build_figure_6()
    copy_unchanged_figures()
    write_manifest()
    print(json.dumps({"status": "BUILT", "target": str(TARGET), "files": len([p for p in TARGET.rglob('*') if p.is_file()])}, ensure_ascii=False))


if __name__ == "__main__":
    main()
